"""Tests for the shared deterministic workflow helpers."""

from __future__ import annotations

import json
import shutil
from contextlib import contextmanager
from pathlib import Path
from uuid import uuid4

from fintools.figures import validate_docx_images_fit_page
from tools import workflow_lib

ROOT = Path(__file__).resolve().parents[1]
TMP_ROOT = ROOT / ".tmp-workflow-tests"


@contextmanager
def temp_repo_dir() -> Path:
    TMP_ROOT.mkdir(exist_ok=True)
    path = TMP_ROOT / f"workflow-{uuid4().hex}"
    path.mkdir(parents=True)
    try:
        yield path
    finally:
        shutil.rmtree(path, ignore_errors=True)


def docx_text(path: Path) -> str:
    """Return text from top-level paragraphs and caption tables."""

    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def run_git(root: Path, *args: str) -> None:
    """Run a git command in a temp repo for workflow tests."""

    report = workflow_lib.run_command(["git", *args], cwd=root, timeout=30)
    assert report.returncode == 0, f"git {' '.join(args)} failed: {report.stderr}"


def make_streamlit_submission_repo(
    root: Path,
    *,
    include_app: bool = True,
    include_config: bool = True,
    app_text: str | None = None,
    tracked_secret: bool = False,
    include_submission_json: bool = True,
    submission_overrides: dict[str, object] | None = None,
) -> Path:
    """Create a minimal clean Git repo with a project Streamlit app."""

    project_root = root / "projects" / "app_project"
    app_root = project_root / "app"
    app_root.mkdir(parents=True)
    (root / ".python-version").write_text("3.13\n", encoding="utf-8")
    (root / "requirements.txt").write_text("streamlit>=1.50,<2\nplotly>=6.1.1\n", encoding="utf-8")
    if include_config:
        (root / ".streamlit").mkdir()
        (root / ".streamlit" / "config.toml").write_text(
            "[browser]\ngatherUsageStats = false\n",
            encoding="utf-8",
        )
    if include_app:
        (app_root / "streamlit_app.py").write_text(
            app_text
            or "\n".join(
                [
                    '"""Test app."""',
                    "import streamlit as st",
                    'st.write("hello")',
                    "",
                ]
            ),
            encoding="utf-8",
        )
    if include_submission_json:
        submission_payload = {
            "repo_url": "https://github.com/example/app-project.git",
            "branch": "main",
            "target": "projects/app_project",
            "entrypoint": "projects/app_project/app/streamlit_app.py",
            "public_app_url": "https://example-app.streamlit.app",
            "repo_visibility": "private",
            "teaching_team_access": {
                "mode": "private_collaborators",
                "github_usernames": ["course-instructor", "course-marker"],
            },
        }
        if submission_overrides:
            submission_payload.update(submission_overrides)
        (project_root / "submission.json").write_text(
            json.dumps(submission_payload, indent=2) + "\n",
            encoding="utf-8",
        )
    if tracked_secret:
        (root / ".env").write_text("TOKEN=secret\n", encoding="utf-8")
    run_git(root, "init")
    run_git(root, "branch", "-M", "main")
    run_git(root, "config", "user.email", "test@example.com")
    run_git(root, "config", "user.name", "Test User")
    run_git(root, "remote", "add", "origin", "https://github.com/example/app-project.git")
    run_git(root, "add", "-A")
    run_git(root, "commit", "-m", "init")
    return project_root


def test_list_command_prints_available_workflows(capsys) -> None:
    status = workflow_lib.main(["list"])
    captured = capsys.readouterr()
    assert status == 0
    assert "audit-week" not in captured.out
    assert "prepare-public-repo" not in captured.out
    assert "new-project" in captured.out
    assert "scaffold-week" in captured.out
    assert "setup-paper" in captured.out
    assert "word-report" in captured.out
    assert "latex-doctor" in captured.out
    assert "build-figure" in captured.out
    assert "build-figure-suite" in captured.out
    assert "build-app" in captured.out
    assert "build-week-context" in captured.out
    assert "check-app-submission" in captured.out
    assert "prepare-app-repo" in captured.out
    assert "section-context" in captured.out
    assert "check-citations" in captured.out


def test_audit_week_passes_for_week1() -> None:
    status, lines = workflow_lib.audit_week(
        ROOT,
        ROOT,
        target="fins2026/week1",
        spec=None,
    )

    assert status == 0
    assert any("Week audit passed." in line for line in lines)


def test_audit_week_passes_for_published_release_weeks() -> None:
    for week_name in ["week2", "week3", "week4", "week5"]:
        status, lines = workflow_lib.audit_week(
            ROOT,
            ROOT,
            target=f"fins2026/{week_name}",
            spec=None,
        )

        assert status == 0, "\n".join(lines)
        assert any("Week audit passed." in line for line in lines)


def test_public_release_filters_blocked_paths() -> None:
    assert workflow_lib.should_exclude_public_release_path(Path(".maintainers/README.md"))
    assert workflow_lib.should_exclude_public_release_path(Path("fins2026/week6/README.md"))
    assert workflow_lib.should_exclude_public_release_path(
        Path("docs/apps/streamlit/validation-rehearsal.md")
    )
    assert workflow_lib.should_exclude_public_release_path(
        Path("docs/apps/streamlit/roadmap.md")
    )
    assert workflow_lib.should_exclude_public_release_path(Path("notes/private.md"))
    assert not workflow_lib.should_exclude_public_release_path(Path("README.md"))
    assert not workflow_lib.should_exclude_public_release_path(
        Path("fins2026/week5/README.md")
    )


def test_prepare_public_repo_creates_sanitized_fresh_history() -> None:
    with temp_repo_dir() as destination:
        status, lines = workflow_lib.prepare_public_repo(
            ROOT,
            ROOT,
            dest=str(destination),
            force=False,
            init_git=True,
        )

        assert status == 0, "\n".join(lines)
        assert not (destination / ".maintainers").exists()
        assert not (destination / "fins2026" / "week6").exists()
        assert not (
            destination / "docs" / "apps" / "streamlit" / "validation-rehearsal.md"
        ).exists()
        assert not (destination / "docs" / "apps" / "streamlit" / "roadmap.md").exists()
        assert (destination / "fins2026" / "week5" / "README.md").exists()
        assert (destination / ".codex" / "config.toml").exists()
        assert (destination / ".github" / "workflows" / "ci.yml").exists()

        audit_status, audit_lines = workflow_lib.audit_public_release_tree(destination)
        assert audit_status == 0, "\n".join(audit_lines)

        tracked_report = workflow_lib.run_command(
            [
                "git",
                "ls-files",
                "fins2026/week5/results/data/stage3/.gitkeep",
                "fins2026/week5/results/tables/stage3/.gitkeep",
            ],
            cwd=destination,
            timeout=30,
        )
        assert tracked_report.returncode == 0, tracked_report.stderr
        tracked_lines = {
            line.strip() for line in tracked_report.stdout.splitlines() if line.strip()
        }
        assert "fins2026/week5/results/data/stage3/.gitkeep" in tracked_lines
        assert "fins2026/week5/results/tables/stage3/.gitkeep" in tracked_lines

        count_report = workflow_lib.run_command(
            ["git", "rev-list", "--count", "HEAD"],
            cwd=destination,
            timeout=30,
        )
        assert count_report.returncode == 0, count_report.stderr
        assert count_report.stdout.strip() == "1"


def test_audit_week_reports_missing_required_file() -> None:
    with temp_repo_dir() as root:
        week_root = root / "fins2026" / "week1"
        (week_root / "README.md").parent.mkdir(parents=True)
        (week_root / "README.md").write_text("# Week 1: Test Week\n", encoding="utf-8")
        spec_dir = root / "internal_specs"
        spec_dir.mkdir(parents=True)
        (spec_dir / "week1.toml").write_text(
            "\n".join(
                [
                    'week_id = "week1"',
                    'title = "Week 1: Test Week"',
                    'required_files = ["README.md", "WORKSHOP.md"]',
                    'results_dirs = ["results/data"]',
                    'script_files = []',
                    'banned_markers = []',
                    "",
                ]
            ),
            encoding="utf-8",
        )

        status, lines = workflow_lib.audit_week(
            root,
            root,
            target="fins2026/week1",
            spec="internal_specs/week1.toml",
        )

        assert status == 1
        assert any("missing required file: WORKSHOP.md" in line for line in lines)


def test_parse_python_version_reads_standard_version_text() -> None:
    assert workflow_lib.parse_python_version("Python 3.13.2") == (3, 13, 2)
    assert workflow_lib.parse_python_version("log\nPython 3.12.9\n") == (3, 12, 9)
    assert workflow_lib.parse_python_version("not python") is None


def test_python_313_candidates_prefer_platform_specific_commands(monkeypatch) -> None:
    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Darwin")
    assert workflow_lib.python_313_candidates()[0][0] == ["python3.13"]

    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Windows")
    assert workflow_lib.python_313_candidates()[0][0] == ["py", "-3.13"]


def test_bootstrap_command_hint_is_platform_specific(monkeypatch) -> None:
    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Windows")
    assert workflow_lib.bootstrap_command_hint() == (
        r"powershell -ExecutionPolicy Bypass -File tools/bootstrap_windows.ps1"
    )

    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Darwin")
    assert workflow_lib.bootstrap_command_hint() == "bash tools/bootstrap_macos.sh"


def test_describe_tabular_file_treats_plain_text_ticker_lists_as_lists() -> None:
    with temp_repo_dir() as root:
        ticker_path = root / "tickers.txt"
        ticker_path.write_text(
            "\n".join(
                [
                    "# Demo ticker universe",
                    "BTC-USD",
                    "ETH-USD",
                    "XRP-USD",
                    "",
                ]
            ),
            encoding="utf-8",
        )

        summary = workflow_lib.describe_tabular_file(ticker_path)

        assert "- Type: `.txt`" in summary
        assert "- Format: plain-text one-item-per-line list" in summary
        assert "- Entries: 3" in summary
        assert any("`BTC-USD`" in line for line in summary)
        assert not any(line.startswith("- Shape:") for line in summary)


def set_onboard_windows(monkeypatch) -> None:
    """Make onboarding decision tests deterministic as Windows/Python 3.13."""

    class VersionInfo:
        major = 3
        minor = 13
        micro = 2

    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Windows")
    monkeypatch.setattr(workflow_lib.sys, "version_info", VersionInfo())
    real_which = shutil.which

    def fake_which(name):
        if name in {"rg", "winget", "brew"}:
            return None
        return real_which(name)

    monkeypatch.setattr(workflow_lib.shutil, "which", fake_which)


def test_onboard_verifies_existing_venv_before_rebuild(monkeypatch) -> None:
    set_onboard_windows(monkeypatch)
    with temp_repo_dir() as root:
        venv_python = root / ".venv" / "Scripts" / "python.exe"
        venv_python.parent.mkdir(parents=True)
        venv_python.write_text("", encoding="utf-8")
        commands: list[list[str]] = []

        def fake_run_command(command, **kwargs):
            del kwargs
            commands.append(command)
            return workflow_lib.CommandReport(command, 0, "setup ok\n", "")

        monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

        status, lines = workflow_lib.onboard(root, root)

        assert status == 0
        assert any("skipping package reinstall" in line for line in lines)
        assert len(commands) == 1
        assert "tools/setup_student.py" in commands[0]


def test_onboard_check_mode_does_not_repair_failed_existing_venv(monkeypatch) -> None:
    set_onboard_windows(monkeypatch)
    with temp_repo_dir() as root:
        venv_python = root / ".venv" / "Scripts" / "python.exe"
        venv_python.parent.mkdir(parents=True)
        venv_python.write_text("", encoding="utf-8")
        commands: list[list[str]] = []

        def fake_run_command(command, **kwargs):
            del kwargs
            commands.append(command)
            return workflow_lib.CommandReport(command, 1, "[BROKEN] pandas\n", "")

        monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

        status, lines = workflow_lib.onboard(root, root, check_only=True)

        assert status == 1
        assert any("Check-only mode requested" in line for line in lines)
        assert len(commands) == 1
        assert "tools/setup_student.py" in commands[0]


def test_onboard_windows_lock_holders_block_repair(monkeypatch) -> None:
    set_onboard_windows(monkeypatch)
    with temp_repo_dir() as root:
        venv_python = root / ".venv" / "Scripts" / "python.exe"
        venv_python.parent.mkdir(parents=True)
        venv_python.write_text("", encoding="utf-8")
        commands: list[list[str]] = []

        def fake_run_command(command, **kwargs):
            del kwargs
            commands.append(command)
            if command[0] == "powershell":
                holder = {
                    "ProcessId": 49568,
                    "Name": "python.exe",
                    "CommandLine": r"C:\repo\.venv\Scripts\python.exe pydevconsole.py",
                }
                return workflow_lib.CommandReport(command, 0, json.dumps([holder]), "")
            return workflow_lib.CommandReport(command, 1, "[BROKEN] pandas\n", "")

        monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

        status, lines = workflow_lib.onboard(root, root)

        assert status == 1
        joined = "\n".join(lines)
        assert "virtual environment is currently used" in joined
        assert "PyCharm Python Console" in joined
        assert not any("-m" in command and "pip" in command for command in commands)


def test_onboard_creates_missing_venv_without_clear(monkeypatch) -> None:
    set_onboard_windows(monkeypatch)
    with temp_repo_dir() as root:
        commands: list[list[str]] = []
        venv_python = root / ".venv" / "Scripts" / "python.exe"

        def fake_run_command(command, **kwargs):
            del kwargs
            commands.append(command)
            if command[1:3] == ["-m", "venv"]:
                assert "--clear" not in command
                venv_python.parent.mkdir(parents=True)
                venv_python.write_text("", encoding="utf-8")
            return workflow_lib.CommandReport(command, 0, "ok\n", "")

        monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

        status, lines = workflow_lib.onboard(root, root)

        assert status == 0
        assert any(command[1:3] == ["-m", "venv"] for command in commands)
        assert not any("--clear" in command for command in commands)
        assert any("All checks passed" in line or "ok" in line for line in lines)


def test_onboard_rebuild_requires_preflight_then_clears_venv(monkeypatch) -> None:
    set_onboard_windows(monkeypatch)
    with temp_repo_dir() as root:
        venv_python = root / ".venv" / "Scripts" / "python.exe"
        venv_python.parent.mkdir(parents=True)
        venv_python.write_text("", encoding="utf-8")
        commands: list[list[str]] = []

        def fake_run_command(command, **kwargs):
            del kwargs
            commands.append(command)
            if command[0] == "powershell":
                return workflow_lib.CommandReport(command, 0, "", "")
            return workflow_lib.CommandReport(command, 0, "ok\n", "")

        monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

        status, lines = workflow_lib.onboard(root, root, rebuild=True)

        assert status == 0
        assert any(command[0] == "powershell" for command in commands)
        assert any(command[1:3] == ["-m", "venv"] and "--clear" in command for command in commands)
        assert any("Rebuild mode requested" in line for line in lines)


def test_onboard_macos_pip_preflight_blocks_pyexpat_failure(monkeypatch) -> None:
    class VersionInfo:
        major = 3
        minor = 13
        micro = 13

    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Darwin")
    monkeypatch.setattr(workflow_lib.sys, "version_info", VersionInfo())
    real_which = shutil.which

    def fake_which(name):
        if name in {"rg", "brew"}:
            return None
        return real_which(name)

    commands: list[list[str]] = []

    def fake_run_command(command, **kwargs):
        del kwargs
        commands.append(command)
        if command == [workflow_lib.sys.executable, "-m", "pip", "--version"]:
            return workflow_lib.CommandReport(
                command,
                1,
                "",
                "ImportError: pyexpat.cpython-313-darwin.so: Symbol not found: "
                "_XML_SetAllocTrackerActivationThreshold Expected in: /usr/lib/libexpat.1.dylib",
            )
        if command[1:3] == ["-m", "venv"]:
            raise AssertionError("venv should not be created after failed pip preflight")
        return workflow_lib.CommandReport(command, 0, "ok\n", "")

    monkeypatch.setattr(workflow_lib.shutil, "which", fake_which)
    monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

    with temp_repo_dir() as root:
        status, lines = workflow_lib.onboard(root, root)

    joined = "\n".join(lines)
    assert status == 1
    assert "Homebrew Python pyexpat/libexpat issue" in joined
    assert "https://www.python.org/downloads/macos/" in joined
    assert "python3.13 -m pip --version" in joined
    assert not any(command[1:3] == ["-m", "venv"] for command in commands)


def test_ripgrep_advisory_check_mode_does_not_install(monkeypatch) -> None:
    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Windows")
    real_which = shutil.which

    def fake_which(name):
        if name == "winget":
            return "winget"
        if name == "rg":
            return None
        return real_which(name)

    def fake_run_command(command, **kwargs):
        del kwargs
        raise AssertionError(f"unexpected command: {command}")

    monkeypatch.setattr(workflow_lib.shutil, "which", fake_which)
    monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

    lines = workflow_lib.ripgrep_advisory(ROOT, install=False)

    joined = "\n".join(lines)
    assert "[WARN] ripgrep" in joined
    assert "winget install --id BurntSushi.ripgrep.MSVC -e" in joined


def test_ripgrep_advisory_windows_attempts_winget_install(monkeypatch) -> None:
    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Windows")
    installed = False
    commands: list[list[str]] = []
    real_which = shutil.which

    def fake_which(name):
        if name == "winget":
            return "winget"
        if name == "rg":
            return "rg" if installed else None
        return real_which(name)

    def fake_run_command(command, **kwargs):
        nonlocal installed
        del kwargs
        commands.append(command)
        if command[0] == "winget":
            installed = True
            return workflow_lib.CommandReport(command, 0, "installed\n", "")
        if command == ["rg", "--version"]:
            return workflow_lib.CommandReport(command, 0, "ripgrep 15.1.0\n", "")
        if command == ["rg", "--files", "-g", "AGENTS.md"]:
            return workflow_lib.CommandReport(command, 0, "AGENTS.md\n", "")
        raise AssertionError(f"unexpected command: {command}")

    monkeypatch.setattr(workflow_lib.shutil, "which", fake_which)
    monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

    lines = workflow_lib.ripgrep_advisory(ROOT, install=True)

    assert any(
        command[:4] == ["winget", "install", "--id", "BurntSushi.ripgrep.MSVC"]
        for command in commands
    )
    assert any("[OK] ripgrep" in line for line in lines)


def test_ripgrep_advisory_macos_attempts_brew_install(monkeypatch) -> None:
    monkeypatch.setattr(workflow_lib.platform, "system", lambda: "Darwin")
    installed = False
    commands: list[list[str]] = []
    real_which = shutil.which

    def fake_which(name):
        if name == "brew":
            return "brew"
        if name == "rg":
            return "rg" if installed else None
        return real_which(name)

    def fake_run_command(command, **kwargs):
        nonlocal installed
        del kwargs
        commands.append(command)
        if command == ["brew", "install", "ripgrep"]:
            installed = True
            return workflow_lib.CommandReport(command, 0, "installed\n", "")
        if command == ["rg", "--version"]:
            return workflow_lib.CommandReport(command, 0, "ripgrep 15.1.0\n", "")
        if command == ["rg", "--files", "-g", "AGENTS.md"]:
            return workflow_lib.CommandReport(command, 0, "AGENTS.md\n", "")
        raise AssertionError(f"unexpected command: {command}")

    monkeypatch.setattr(workflow_lib.shutil, "which", fake_which)
    monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

    lines = workflow_lib.ripgrep_advisory(ROOT, install=True)

    assert ["brew", "install", "ripgrep"] in commands
    assert any("[OK] ripgrep" in line for line in lines)


def test_onboard_continues_when_ripgrep_install_fails(monkeypatch) -> None:
    set_onboard_windows(monkeypatch)
    with temp_repo_dir() as root:
        commands: list[list[str]] = []
        venv_python = root / ".venv" / "Scripts" / "python.exe"
        real_which = shutil.which

        def fake_which(name):
            if name == "winget":
                return "winget"
            if name == "rg":
                return None
            return real_which(name)

        def fake_run_command(command, **kwargs):
            del kwargs
            commands.append(command)
            if command[0] == "winget":
                return workflow_lib.CommandReport(command, 1, "", "winget failed\n")
            if command[1:3] == ["-m", "venv"]:
                venv_python.parent.mkdir(parents=True)
                venv_python.write_text("", encoding="utf-8")
            return workflow_lib.CommandReport(command, 0, "ok\n", "")

        monkeypatch.setattr(workflow_lib.shutil, "which", fake_which)
        monkeypatch.setattr(workflow_lib, "run_command", fake_run_command)

        status, lines = workflow_lib.onboard(root, root)

        assert status == 0
        assert any(command[0] == "winget" for command in commands)
        assert any(command[1:3] == ["-m", "venv"] for command in commands)
        assert any("Could not install ripgrep automatically" in line for line in lines)


def test_create_project_scaffold_creates_expected_files() -> None:
    with temp_repo_dir() as tmp_path:
        (tmp_path / "projects").mkdir()
        project_root, lines = workflow_lib.create_project(
            tmp_path,
            name="alpha_project",
            description="Event study on bank earnings surprises.",
            datasets="WRDS CRSP monthly returns",
            notes="Keep the final report under report/.",
        )
        assert project_root == tmp_path / "projects" / "alpha_project"
        assert (project_root / "report").is_dir()
        assert (project_root / "guidance").is_dir()
        assert (project_root / "README.md").exists()
        assert (project_root / "AGENTS.md").exists()
        assert (project_root / "scripts" / "make_figures.py").exists()
        readme = (project_root / "README.md").read_text(encoding="utf-8")
        assert "python ../../tools/workflow.py setup-paper" in readme
        assert "report/report.docx" in readme
        assert "python scripts/make_figures.py" in readme
        assert "Created project scaffold" in lines[0]


def test_create_project_scaffold_can_include_streamlit_app() -> None:
    with temp_repo_dir() as tmp_path:
        (tmp_path / "projects").mkdir()
        project_root, lines = workflow_lib.create_project(
            tmp_path,
            name="app_project",
            description="Forecast macro fundamentals.",
            datasets="FRED",
            notes="Deploy on Streamlit Community Cloud.",
            with_app=True,
        )
        app_path = project_root / "app" / "streamlit_app.py"
        assert app_path.exists()
        assert (project_root / ".streamlit" / "config.toml").exists()
        checklist = project_root / "SUBMISSION_CHECKLIST.md"
        assert checklist.exists()
        assert not (project_root / "submission.json").exists()
        assert (project_root / "app" / "tests" / "test_app_smoke.py").exists()
        assert (project_root / "results" / "app").is_dir()
        app_text = app_path.read_text(encoding="utf-8")
        checklist_text = checklist.read_text(encoding="utf-8")
        assert "streamlit" in app_text
        assert "sys.path.insert" in app_text
        assert "from fintools.apps import" in app_text
        assert app_text.index("sys.path.insert") < app_text.index("from fintools.apps import")
        assert "st.segmented_control" in app_text
        assert "Sample period" in app_text
        assert "lazy_tabs" in app_text
        assert "query_choice" in app_text
        assert "sync_query_params" in app_text
        assert "render_data_health" in app_text
        assert "render_display_table" in app_text
        assert "render_csv_download" in app_text
        assert "target_forecast_figure" in app_text
        assert "forecast_series_spec" in app_text
        assert "SeriesSpec" in app_text
        assert "Absolute error" in app_text
        assert "C:\\Users" not in app_text
        assert ".maintainers" not in app_text
        assert "Public Streamlit app URL" in checklist_text
        assert "Final commit hash" in checklist_text
        assert "Required Deployment Fields" in checklist_text
        assert any("Streamlit app scaffold" in line for line in lines)


def test_scaffold_week_creates_expected_files() -> None:
    with temp_repo_dir() as root:
        (root / "fins2026" / "week4").mkdir(parents=True)
        week_root, lines = workflow_lib.scaffold_week(root, root, target="fins2026/week4")
        assert week_root == root / "fins2026" / "week4"
        assert (week_root / "README.md").exists()
        assert (week_root / "WORKSHOP.md").exists()
        assert (week_root / "DATA_GUIDE.md").exists()
        assert (week_root / "AGENTS.md").exists()
        assert (week_root / "guidance" / "week-context.md").exists()
        assert (week_root / "guidance" / "data-context.md").exists()
        assert (week_root / "guidance" / "output-context.md").exists()
        assert (week_root / "scripts" / "run_week.py").exists()
        assert (week_root / "scripts" / "describe_data.py").exists()
        assert (week_root / "tests" / "test_week_smoke.py").exists()
        assert (week_root / "results" / "figures" / ".gitkeep").exists()
        assert (week_root / "results" / "data" / ".gitkeep").exists()
        readme = (week_root / "README.md").read_text(encoding="utf-8")
        assert "build-week-context" in readme
        assert "scratch/" in readme
        assert any("Generated week context" in line for line in lines)


def test_scaffold_week_preserves_existing_authored_files() -> None:
    with temp_repo_dir() as root:
        week_root = root / "fins2026" / "week6"
        week_root.mkdir(parents=True)
        (week_root / "README.md").write_text("# Custom Week 6\n\nAuthor text.\n", encoding="utf-8")

        workflow_lib.scaffold_week(root, root, target="fins2026/week6")

        assert (week_root / "README.md").read_text(encoding="utf-8") == (
            "# Custom Week 6\n\nAuthor text.\n"
        )
        assert (week_root / "guidance" / "week-context.md").exists()
        week_context = (week_root / "guidance" / "week-context.md").read_text(encoding="utf-8")
        assert "Custom Week 6" in week_context


def test_build_week_context_summarizes_week_data_and_outputs() -> None:
    with temp_repo_dir() as root:
        week_root = root / "fins2026" / "week8"
        week_root.mkdir(parents=True)
        workflow_lib.scaffold_week(root, root, target="fins2026/week8")
        (week_root / "APP_LAB.md").write_text(
            "# App Lab\n\nTurn this week into an app.\n",
            encoding="utf-8",
        )
        (week_root / "data" / "demo.csv").write_text(
            "date,value,segment\n2024-01-31,1.0,IG\n2024-02-29,1.5,HY\n",
            encoding="utf-8",
        )
        (week_root / "results" / "data" / "derived.csv").write_text(
            "date,signal\n2024-02-29,2.5\n",
            encoding="utf-8",
        )
        (week_root / "results" / "figures" / "chart.png").write_text("png\n", encoding="utf-8")
        (week_root / "results" / "forecasts").mkdir(parents=True, exist_ok=True)
        (week_root / "results" / "forecasts" / "leaderboard.csv").write_text(
            "model,target_rmse\nnaive,1.0\n",
            encoding="utf-8",
        )

        output_paths, _ = workflow_lib.build_week_context(root, root, target="fins2026/week8")

        assert len(output_paths) == 3
        data_context = (week_root / "guidance" / "data-context.md").read_text(encoding="utf-8")
        output_context = (week_root / "guidance" / "output-context.md").read_text(
            encoding="utf-8"
        )
        week_context = (week_root / "guidance" / "week-context.md").read_text(encoding="utf-8")
        assert "demo.csv" in data_context
        assert "Shape: 2 rows x 3 columns" in data_context
        assert "`date`" in data_context
        assert "chart.png" in output_context
        assert "Figure Outputs" in output_context
        assert "Forecast Outputs" in output_context
        assert "leaderboard.csv" in output_context
        assert "App Lab" in week_context


def test_build_week_context_ignores_untracked_generated_outputs_in_git_repo() -> None:
    with temp_repo_dir() as root:
        week_root = root / "fins2026" / "week8"
        workflow_lib.scaffold_week(root, root, target="fins2026/week8")
        (week_root / "data" / "demo.csv").write_text(
            "date,value\n2024-01-31,1.0\n",
            encoding="utf-8",
        )
        run_git(root, "init")
        run_git(root, "branch", "-M", "main")
        run_git(root, "config", "user.email", "test@example.com")
        run_git(root, "config", "user.name", "Test User")
        run_git(root, "add", "-A")
        run_git(root, "commit", "-m", "init")

        (week_root / "results" / "data" / "derived.csv").write_text(
            "date,signal\n2024-02-29,2.5\n",
            encoding="utf-8",
        )
        (week_root / "results" / "figures" / "chart.png").write_text("png\n", encoding="utf-8")
        (week_root / "results" / "forecasts").mkdir(parents=True, exist_ok=True)
        (week_root / "results" / "forecasts" / "leaderboard.csv").write_text(
            "model,target_rmse\nnaive,1.0\n",
            encoding="utf-8",
        )

        workflow_lib.build_week_context(root, root, target="fins2026/week8")

        data_context = (week_root / "guidance" / "data-context.md").read_text(encoding="utf-8")
        output_context = (week_root / "guidance" / "output-context.md").read_text(
            encoding="utf-8"
        )
        assert "demo.csv" in data_context
        assert "derived.csv" not in data_context
        assert "chart.png" not in output_context
        assert "leaderboard.csv" not in output_context
        assert "generated locally and not committed by default" in output_context


def test_build_app_scaffold_creates_app_in_existing_folder() -> None:
    with temp_repo_dir() as root:
        target = root / "fins2026" / "week9"
        target.mkdir(parents=True)
        workflow_lib.scaffold_week(root, root, target="fins2026/week9")
        app_root, lines = workflow_lib.build_app_scaffold(
            root,
            root,
            target="fins2026/week9",
            title="Week 9 Credit App",
            description="Credit-market insight app.",
        )
        assert app_root == target / "app"
        assert (app_root / "streamlit_app.py").exists()
        assert (target / "SUBMISSION_CHECKLIST.md").exists()
        assert not (target / "submission.json").exists()
        assert (target / "results" / "app").is_dir()
        assert "streamlit run fins2026/week9/app/streamlit_app.py" in "\n".join(lines)
        checklist_text = (target / "SUBMISSION_CHECKLIST.md").read_text(encoding="utf-8")
        assert "fins2026/week9/app/streamlit_app.py" in checklist_text
        assert "Required Deployment Fields" in checklist_text


def test_check_app_submission_reports_ready_for_clean_project_repo() -> None:
    with temp_repo_dir() as root:
        project_root = make_streamlit_submission_repo(root)
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        text = "\n".join(lines)
        assert status == 0
        assert "Status: READY" in text
        assert "Deployment metadata:" in text
        assert "Teaching-team access: private_collaborators" in text
        assert "Repository: https://github.com/example/app-project.git" in text
        assert "Branch: main" in text
        assert "Entrypoint: projects/app_project/app/streamlit_app.py" in text
        assert "Public Streamlit app URL: https://example-app.streamlit.app" in text
        assert project_root.name in text


def test_check_app_submission_uses_inferred_repo_state_without_submission_json() -> None:
    with temp_repo_dir() as root:
        make_streamlit_submission_repo(root, include_submission_json=False)
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        text = "\n".join(lines)
        assert status == 0
        assert "Status: READY" in text
        assert "Deployment metadata:" in text
        assert "- File: not present" in text
        assert "- Mode: inferred from the target, entrypoint, and Git state" in text


def test_check_app_submission_blocks_mismatched_submission_entrypoint() -> None:
    with temp_repo_dir() as root:
        make_streamlit_submission_repo(
            root,
            submission_overrides={"entrypoint": "projects/app_project/app/not_the_app.py"},
        )
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        assert status == 1
        assert "submission.json entrypoint does not match" in "\n".join(lines)


def test_check_app_submission_blocks_missing_entrypoint() -> None:
    with temp_repo_dir() as root:
        make_streamlit_submission_repo(root, include_app=False)
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        assert status == 1
        assert "Streamlit entrypoint does not exist" in "\n".join(lines)


def test_check_app_submission_blocks_missing_root_streamlit_config() -> None:
    with temp_repo_dir() as root:
        make_streamlit_submission_repo(root, include_config=False)
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        assert status == 1
        assert "Missing root .streamlit/config.toml" in "\n".join(lines)


def test_check_app_submission_blocks_tracked_secret_files() -> None:
    with temp_repo_dir() as root:
        make_streamlit_submission_repo(root, tracked_secret=True)
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        text = "\n".join(lines)
        assert status == 1
        assert "Tracked secret/environment files found" in text
        assert ".env" in text


def test_check_app_submission_blocks_local_absolute_paths() -> None:
    with temp_repo_dir() as root:
        make_streamlit_submission_repo(
            root,
            app_text="\n".join(
                [
                    '"""Test app."""',
                    "import streamlit as st",
                    'LOCAL = r"C:\\Users\\alexm\\Desktop\\data.csv"',
                    "st.write(LOCAL)",
                    "",
                ]
            ),
        )
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        assert status == 1
        assert "Local absolute paths found under app target" in "\n".join(lines)


def test_check_app_submission_blocks_dirty_working_tree() -> None:
    with temp_repo_dir() as root:
        project_root = make_streamlit_submission_repo(root)
        (project_root / "README.md").write_text("uncommitted change\n", encoding="utf-8")
        status, lines = workflow_lib.check_app_submission(
            root,
            root,
            target="projects/app_project",
        )
        assert status == 1
        assert "Working tree is not clean" in "\n".join(lines)


def test_check_app_submission_week2_reports_exact_deploy_fields() -> None:
    status, lines = workflow_lib.check_app_submission(
        ROOT,
        ROOT,
        target="fins2026/week2",
        entrypoint="fins2026/week2/app/streamlit_app.py",
    )
    text = "\n".join(lines)
    assert status in {0, 1}
    assert "Entrypoint: fins2026/week2/app/streamlit_app.py" in text
    assert "Main file path: fins2026/week2/app/streamlit_app.py" in text
    assert "Python version: 3.13" in text
    assert "Secrets: none detected" in text


def test_prepare_app_repo_creates_minimal_deployable_week2_bundle() -> None:
    with temp_repo_dir() as tmp_path:
        destination = tmp_path / "week2-handin"
        status, lines = workflow_lib.prepare_app_repo(
            ROOT,
            ROOT,
            source="fins2026/week2",
            dest=str(destination),
            repo="week2-streamlit-handin-rehearsal",
            entrypoint="fins2026/week2/app/streamlit_app.py",
        )
        text = "\n".join(lines)
        assert status == 0
        assert "Prepared Streamlit app deploy repository" in text
        assert "Status: WARNINGS" in text
        assert "Git remote 'origin' is not configured yet" in text
        assert (destination / ".git").is_dir()
        assert (destination / "requirements.txt").exists()
        assert (destination / "requirements-dev.txt").exists()
        requirements_dev_text = (destination / "requirements-dev.txt").read_text(
            encoding="utf-8"
        )
        assert "-e ." not in requirements_dev_text
        assert (destination / ".python-version").exists()
        assert (destination / ".streamlit" / "config.toml").exists()
        assert (destination / ".github" / "workflows" / "submission-check.yml").exists()
        assert (destination / "tools" / "workflow.py").exists()
        assert (destination / "tools" / "workflow_lib.py").exists()
        assert (destination / "fintools" / "apps" / "forecasting.py").exists()
        assert (destination / "fintools" / "datasets" / "validation").is_dir()
        assert (
            destination / "fins2026" / "week2" / "app" / "streamlit_app.py"
        ).exists()
        assert (destination / "README.md").exists()
        assert (destination / "SUBMISSION_CHECKLIST.md").exists()
        assert not (destination / "submission.json").exists()
        assert not (destination / ".venv").exists()
        assert not (destination / ".maintainers").exists()
        assert not (destination / "fintools" / "datasets" / "validation" / "_raw").exists()
        assert not (destination / "fins2026" / "week2" / "results").exists()
        workflow_text = (
            destination / ".github" / "workflows" / "submission-check.yml"
        ).read_text(encoding="utf-8")
        assert "--target fins2026/week2" in workflow_text
        assert "--entrypoint fins2026/week2/app/streamlit_app.py" in workflow_text

        ready_status, ready_lines = workflow_lib.check_app_submission(
            destination,
            destination,
            target="fins2026/week2",
            entrypoint="fins2026/week2/app/streamlit_app.py",
            require_remote=False,
        )
        assert ready_status == 0
        assert "Main file path: fins2026/week2/app/streamlit_app.py" in "\n".join(
            ready_lines
        )


def test_prepare_app_repo_refuses_non_empty_destination_without_force() -> None:
    with temp_repo_dir() as tmp_path:
        destination = tmp_path / "week2-handin"
        destination.mkdir()
        (destination / "keep.txt").write_text("do not delete\n", encoding="utf-8")
        try:
            workflow_lib.prepare_app_repo(
                ROOT,
                ROOT,
                source="fins2026/week2",
                dest=str(destination),
                repo="week2-streamlit-handin-rehearsal",
                entrypoint="fins2026/week2/app/streamlit_app.py",
            )
        except workflow_lib.WorkflowError as exc:
            assert "destination is not empty" in str(exc)
        else:  # pragma: no cover - defensive
            raise AssertionError("expected non-empty destination to fail")


def test_prepare_app_repo_push_reports_missing_github_cli_auth(monkeypatch) -> None:
    with temp_repo_dir() as tmp_path:
        destination = tmp_path / "week2-handin"
        original_which = shutil.which

        def fake_which(name: str) -> str | None:
            if name == "gh":
                return None
            return original_which(name)

        monkeypatch.setattr(workflow_lib.shutil, "which", fake_which)
        status, lines = workflow_lib.prepare_app_repo(
            ROOT,
            ROOT,
            source="fins2026/week2",
            dest=str(destination),
            repo="week2-streamlit-handin-rehearsal",
            entrypoint="fins2026/week2/app/streamlit_app.py",
            push=True,
        )
        text = "\n".join(lines)
        assert status == 1
        assert "GitHub push skipped: GitHub CLI is not installed." in text
        assert (destination / ".git").is_dir()


def test_build_figure_examples_writes_validation_outputs() -> None:
    with temp_repo_dir() as root:
        status, lines = workflow_lib.build_figure_examples(
            root,
            root,
            output="results/figures",
            docx=True,
        )
        assert status == 0
        assert (root / "results" / "figures" / "validation_market_cumulative.png").exists()
        assert (root / "results" / "figures" / "validation_ff3_full_returns.png").exists()
        assert (root / "results" / "figures" / "validation_ff3_full_cumulative.png").exists()
        assert (root / "results" / "figures" / "validation_macro_indexed.png").exists()
        assert (root / "results" / "figures" / "validation_shiller_real_market.png").exists()
        assert (root / "results" / "figures" / "validation_macro_scatter_fit.png").exists()
        assert (root / "results" / "figures" / "validation_industry_small_multiples.png").exists()
        assert (root / "results" / "figures" / "validation_ff25_mean_heatmap.png").exists()
        assert (root / "results" / "figures" / "validation_world_bank_bubble.png").exists()
        assert (root / "results" / "figures" / "validation_stress_calendar_heatmap.png").exists()
        assert (root / "results" / "figures" / "validation_figures.docx").exists()
        assert not validate_docx_images_fit_page(
            root / "results" / "figures" / "validation_figures.docx"
        )
        proof_text = docx_text(root / "results" / "figures" / "validation_figures.docx")
        assert "Figure 1. Full-Sample Fama/French Factor Returns." in proof_text
        assert "Figure 14. Selected Industry Return Small Multiples." in proof_text
        assert "Average Returns Across Size-Value Portfolios" in proof_text
        assert "Population, Income, And GDP Scale" in proof_text
        assert "Daily VIX Calendar Heatmap During 2020" in proof_text
        assert "The sample spans the time period" in proof_text
        assert not (root / "results" / "figures" / "validation_market_cumulative.docx").exists()
        assert (
            not (root / "results" / "figures" / "validation_factor_stacked_returns.docx").exists()
        )
        assert not (root / "results" / "figures" / "validation_industry_mean_returns.docx").exists()
        assert not (root / "results" / "figures" / "validation_industry_correlations.docx").exists()
        assert not (root / "results" / "figures" / "validation_macro_rates.docx").exists()
        assert (
            root / "results" / "figures" / "validation_industry_correlations.pdf"
        ).exists()
        assert any("validation_figures.docx" in line for line in lines)
        assert any("validation_macro_scatter_fit.caption.md" in line for line in lines)


def test_build_figure_examples_writes_ft_validation_outputs() -> None:
    with temp_repo_dir() as root:
        status, lines = workflow_lib.build_figure_examples(
            root,
            root,
            output="results/figures-ft",
            docx=True,
            style="ft",
        )
        assert status == 0
        assert (root / "results" / "figures-ft" / "ft_validation_ff3_full_returns.png").exists()
        assert (
            root / "results" / "figures-ft" / "ft_validation_world_bank_gdp_lollipop.png"
        ).exists()
        assert (
            root / "results" / "figures-ft" / "ft_validation_macro_policy_episode.png"
        ).exists()
        assert (
            root / "results" / "figures-ft" / "ft_validation_ff25_mean_heatmap.png"
        ).exists()
        assert (
            root / "results" / "figures-ft" / "ft_validation_world_bank_bubble.png"
        ).exists()
        assert (
            root / "results" / "figures-ft" / "ft_validation_stress_calendar_heatmap.png"
        ).exists()
        assert (root / "results" / "figures-ft" / "validation_figures_ft.docx").exists()
        assert not validate_docx_images_fit_page(
            root / "results" / "figures-ft" / "validation_figures_ft.docx"
        )
        proof_text = docx_text(root / "results" / "figures-ft" / "validation_figures_ft.docx")
        assert "FT-Style Validation Figure Examples" in proof_text
        assert "Largest Economies In Current U.S. Dollars" in proof_text
        assert "other economies are muted comparison points" in proof_text
        assert "Policy Rates And Unemployment During The Financial Crisis" in proof_text
        assert "Average Returns Across Size-Value Portfolios" in proof_text
        assert "Yield-Curve Inversions Against Zero" in proof_text
        assert any("validation_figures_ft.docx" in line for line in lines)


def test_build_figure_suite_from_csv_writes_dataframe_outputs() -> None:
    with temp_repo_dir() as root:
        data_dir = root / "data"
        data_dir.mkdir()
        csv_path = data_dir / "demo.csv"
        csv_path.write_text(
            "date,market_return,smb_return,credit_spread_percent,volume_usd,segment\n"
            "2020-01-31,1.1,0.2,3.1,100,IG\n"
            "2020-02-29,-2.0,0.4,3.4,105,HY\n"
            "2020-03-31,-8.0,-1.2,5.2,90,IG\n"
            "2020-04-30,5.5,1.1,4.8,110,HY\n"
            "2020-05-31,2.2,0.3,4.0,116,IG\n"
            "2020-06-30,1.8,-0.2,3.8,120,HY\n"
            "2020-07-31,0.9,0.1,3.5,124,IG\n"
            "2020-08-31,1.3,0.2,3.3,128,HY\n"
            "2020-09-30,-0.4,-0.1,3.6,122,IG\n"
            "2020-10-31,1.0,0.0,3.2,130,HY\n"
            "2020-11-30,3.2,0.6,3.0,136,IG\n"
            "2020-12-31,2.4,0.5,2.9,140,HY\n",
            encoding="utf-8",
        )
        status, lines = workflow_lib.build_figure_suite(
            root,
            root,
            input_path="data/demo.csv",
            output="results/figures-suite",
            docx=False,
            style="ft",
            title_prefix="Demo",
            max_figures=3,
        )
        assert status == 0
        figure_dir = root / "results" / "figures-suite"
        assert (figure_dir / "figure_suite_01_wide_time_series.png").exists()
        assert (figure_dir / "figure_suite_01_wide_time_series.pdf").exists()
        assert not (root / "results" / "figures-suite" / "figure_suite.docx").exists()
        assert any("Profile:" in line for line in lines)
        assert any("wide_time_series" in line for line in lines)


def test_setup_paper_creates_word_report_by_default() -> None:
    with temp_repo_dir() as root:
        project_dir = root / "projects" / "alpha_project"
        (project_dir / "report").mkdir(parents=True)

        target_dir, lines = workflow_lib.setup_paper(
            root,
            project_dir,
            title="Alpha Project Results",
            authors="Student Name; z1234567",
            topic="bank earnings surprises",
            target=".",
        )
        report_path = target_dir / "report.docx"
        text = docx_text(report_path)
        assert target_dir == project_dir / "report"
        assert report_path.exists()
        assert "Alpha Project Results" in text
        assert "Student Name; z1234567" in text
        assert "bank earnings surprises" in text
        assert any("report.docx" in line for line in lines)


def test_setup_paper_can_create_legacy_latex_scaffold() -> None:
    with temp_repo_dir() as root:
        boilerplate = root / "boilerplate"
        boilerplate.mkdir()
        (boilerplate / "template_main.tex").write_text(
            "\\title{[REMOVE] Your Report Title Here}\n"
            "\\author{[REMOVE] Student Name \\\\ [REMOVE] Student ID \\\\ FINS2026 --- Fintech}\n"
            "\\begin{abstract}\n"
            "[REMOVE] Write a concise summary of your report here. State the research\n"
            "question, the data or methodology used, and the main findings. Aim for\n"
            "100--200 words.\n"
            "\\end{abstract}\n",
            encoding="utf-8",
        )
        (boilerplate / "template_references.bib").write_text(
            "@article{Demo_2024,}\n",
            encoding="utf-8",
        )
        project_dir = root / "projects" / "alpha_project"
        (project_dir / "latex").mkdir(parents=True)

        target_dir, lines = workflow_lib.setup_paper(
            root,
            project_dir,
            title="Alpha Project Results",
            authors="Student Name; z1234567",
            topic="bank earnings surprises",
            target=".",
            format="latex",
        )
        text = (target_dir / "main.tex").read_text(encoding="utf-8")
        assert target_dir == project_dir / "latex"
        assert "\\title{Alpha Project Results}" in text
        assert "Student Name \\\\ z1234567" in text
        assert "bank earnings surprises" in text
        assert any("main.tex" in line for line in lines)


def test_build_context_writes_guidance_file_from_tex() -> None:
    with temp_repo_dir() as root:
        (root / "projects").mkdir()
        project_dir = root / "projects" / "alpha_project"
        latex_dir = project_dir / "latex"
        latex_dir.mkdir(parents=True)
        tex_path = latex_dir / "main.tex"
        tex_path.write_text(
            "\\title{Alpha Project}\n"
            "\\begin{abstract}This report studies market microstructure.\\end{abstract}\n"
            "%% BEGIN data\n"
            "\\section{Data}\\label{sec:data}\n"
            "We use TAQ data from 2015 to 2024.\n"
            "%% END data\n"
            "%% BEGIN results\n"
            "\\section{Results}\\label{sec:results}\n"
            "The main result is a 12 basis point return spread.\n"
            "%% END results\n"
            "\\citep{Demo_2024}\n",
            encoding="utf-8",
        )

        output_path, lines = workflow_lib.build_context(root, latex_dir, sources=["main.tex"])
        text = output_path.read_text(encoding="utf-8")
        assert output_path == project_dir / "guidance" / "paper-context.md"
        assert "Alpha Project" in text
        assert "TAQ data" in text
        assert "Demo_2024" in text
        assert any("paper-context.md" in line for line in lines)


def test_build_context_writes_guidance_file_from_word_report() -> None:
    with temp_repo_dir() as root:
        project_dir = root / "projects" / "alpha_project"
        (project_dir / "report").mkdir(parents=True)
        _, _ = workflow_lib.setup_paper(
            root,
            project_dir,
            title="Alpha Project Results",
            authors="Student Name",
            topic="market microstructure",
            target=".",
        )

        output_path, lines = workflow_lib.build_context(
            root,
            project_dir,
            sources=["report/report.docx"],
        )
        text = output_path.read_text(encoding="utf-8")
        assert output_path == project_dir / "guidance" / "paper-context.md"
        assert "Alpha Project Results" in text
        assert "Abstract" in text
        assert "Introduction" in text
        assert any("paper-context.md" in line for line in lines)


def test_outline_report_detects_missing_standard_sections() -> None:
    with temp_repo_dir() as root:
        tex_path = root / "main.tex"
        tex_path.write_text(
            "%% BEGIN introduction\n"
            "\\section{Introduction}\n"
            "Short intro.\n"
            "%% END introduction\n"
            "%% BEGIN results\n"
            "\\section{Results}\n"
            "A much longer results section.\n"
            "Line two.\n"
            "Line three.\n"
            "%% END results\n",
            encoding="utf-8",
        )

        status, lines = workflow_lib.outline_report(root, root, target="main.tex")
        assert status == 0
        joined = "\n".join(lines)
        assert "Missing standard sections" in joined
        assert "Marker spans" in joined


def test_outline_report_reads_word_headings() -> None:
    with temp_repo_dir() as root:
        project_dir = root / "projects" / "alpha_project"
        (project_dir / "report").mkdir(parents=True)
        workflow_lib.setup_paper(
            root,
            project_dir,
            title="Alpha Project Results",
            target=".",
        )

        status, lines = workflow_lib.outline_report(
            root,
            project_dir,
            target="report/report.docx",
        )
        assert status == 0
        joined = "\n".join(lines)
        assert "Headings:" in joined
        assert "Introduction" in joined
        assert "Marker spans: not used for Word reports." in joined


def test_proofread_report_finds_mechanical_issues() -> None:
    with temp_repo_dir() as root:
        tex_path = root / "main.tex"
        tex_path.write_text(
            "%% BEGIN intro\n"
            "\\section{Introduction}\n"
            "This this sentence has a doubled word.\n"
            "Figure \\ref{fig:test} shows the result.\n"
            "[REMOVE] Fix me later.\n"
            "%% END intro\n",
            encoding="utf-8",
        )

        status, lines = workflow_lib.proofread_report(
            root,
            root,
            target="main.tex",
            section_key="intro",
            line_range=None,
        )
        assert status == 1
        joined = "\n".join(lines)
        assert "doubled word" in joined
        assert "nonbreaking space" in joined
        assert "[REMOVE]" in joined


def test_proofread_report_reads_word_paragraphs() -> None:
    with temp_repo_dir() as root:
        from docx import Document

        docx_path = root / "report.docx"
        document = Document()
        document.add_heading("Introduction", level=1)
        document.add_paragraph("This this paragraph has a doubled word.")
        document.add_paragraph("[REMOVE] Fix me later.")
        document.save(docx_path)

        status, lines = workflow_lib.proofread_report(
            root,
            root,
            target="report.docx",
            section_key="introduction",
            line_range=None,
        )
        assert status == 1
        joined = "\n".join(lines)
        assert "paragraph" in joined
        assert "doubled word" in joined
        assert "[REMOVE]" in joined


def test_latex_doctor_reports_comment_candidates_and_marker_issues() -> None:
    with temp_repo_dir() as root:
        tex_path = root / "main.tex"
        tex_path.write_text(
            "% ordinary comment\n"
            "%% BEGIN intro\n"
            "\\section{Introduction}\n"
            "Content.\n"
            "%% END wrong_key\n",
            encoding="utf-8",
        )

        status, lines = workflow_lib.latex_doctor(root, root, target="main.tex", mode="markers")
        assert status == 1
        assert any("Marker issues:" in line for line in lines)

        status, lines = workflow_lib.latex_doctor(root, root, target="main.tex", mode="comments")
        assert status == 0
        assert any("Comment cleanup candidates" in line for line in lines)


def test_section_context_reports_labels_and_citations() -> None:
    with temp_repo_dir() as root:
        tex_path = root / "main.tex"
        tex_path.write_text(
            "%% BEGIN results\n"
            "\\section{Results}\\label{sec:results}\n"
            "We build on \\citet{Demo_2024}.\n"
            "%% END results\n",
            encoding="utf-8",
        )

        status, lines = workflow_lib.section_context_report(
            root,
            root,
            target="main.tex",
            section_key="results",
        )
        assert status == 0
        joined = "\n".join(lines)
        assert "sec:results" in joined
        assert "Demo_2024" in joined


def test_check_citations_flags_missing_keys() -> None:
    with temp_repo_dir() as root:
        tex_path = root / "main.tex"
        tex_path.write_text(
            "\\bibliography{references}\n"
            "%% BEGIN intro\n"
            "\\section{Introduction}\n"
            "As shown by \\citet{Known_2024} and \\citet{Missing_2025}.\n"
            "%% END intro\n",
            encoding="utf-8",
        )
        (root / "references.bib").write_text(
            "@article{Known_2024,\n  title = {Known paper}\n}\n",
            encoding="utf-8",
        )

        status, lines = workflow_lib.check_citations_report(
            root,
            root,
            target="main.tex",
            section_key="intro",
            line_range=None,
        )
        assert status == 1
        joined = "\n".join(lines)
        assert "Known_2024" in joined
        assert "Missing_2025" in joined
