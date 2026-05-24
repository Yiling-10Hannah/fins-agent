"""Tests for publication-quality figure helpers."""

from __future__ import annotations

import shutil
from contextlib import contextmanager
from pathlib import Path
from uuid import uuid4

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import pytest

from fintools.datasets import available_validation_datasets, load_validation_dataset
from fintools.figures import (
    FT_BACKGROUND,
    FigureContext,
    WordFigureEntry,
    add_nber_recession_shading,
    add_source_note,
    area_balance_plot,
    bar_plot,
    bubble_matrix_plot,
    bubble_scatter_plot,
    calendar_heatmap,
    connected_scatter_plot,
    correlation_heatmap,
    create_figure_suite,
    cumulative_returns_plot,
    distribution_comparison_plot,
    distribution_plot,
    diverging_bar_plot,
    drawdown_plot,
    dumbbell_plot,
    ecdf_plot,
    export_figure_bundle,
    export_word_figure,
    grouped_bar_plot,
    indexed_time_series_plot,
    infer_return_scale,
    insert_figure_docx,
    insert_figures_docx,
    lollipop_plot,
    mean_return_bar_plot,
    plan_figure_suite,
    profile_dataframe,
    proportional_stacked_bar_plot,
    recession_windows_for_range,
    rolling_stat_plot,
    scatter_plot,
    slope_chart,
    small_multiples,
    stacked_area_plot,
    stacked_bar_plot,
    theme_rc,
    time_series_plot,
    uncertainty_band_plot,
    validate_axes_labels,
    validate_category_label_count,
    validate_display_labels,
    validate_docx_images_fit_page,
    validate_equal_subplot_widths,
    validate_figure_context,
    validate_horizontal_grid,
    validate_image_not_blank,
    validate_legend_present,
    validate_markers_within_axes,
    validate_no_text_overlap,
    validate_no_tick_label_overlap,
    validate_series_identification,
    validate_titles_within_canvas,
    validate_unique_series_colors,
    validate_word_readability,
    value_heatmap,
)

ROOT = Path(__file__).resolve().parents[1]
TMP_ROOT = ROOT / ".tmp-figure-tests"


def close(fig) -> None:
    plt.close(fig)


def docx_text(path: Path) -> str:
    """Return text from top-level paragraphs and caption tables."""

    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


@contextmanager
def temp_figure_dir() -> Path:
    TMP_ROOT.mkdir(exist_ok=True)
    path = TMP_ROOT / f"figures-{uuid4().hex}"
    path.mkdir(parents=True)
    try:
        yield path
    finally:
        shutil.rmtree(path, ignore_errors=True)


def test_dataframe_figure_suite_profiles_plans_and_exports() -> None:
    rng = np.random.default_rng(123)
    dates = pd.date_range("2018-01-31", periods=72, freq="ME")
    frame = pd.DataFrame(
        {
            "date": dates,
            "market_return": rng.normal(0.7, 3.8, len(dates)),
            "smb_return": rng.normal(0.1, 1.7, len(dates)),
            "credit_spread_percent": rng.normal(3.0, 0.35, len(dates)).clip(0.5),
            "volume_usd": np.linspace(100, 165, len(dates)) + rng.normal(0, 4, len(dates)),
            "segment": np.where(np.arange(len(dates)) % 2 == 0, "IG", "HY"),
        }
    )

    profile = profile_dataframe(frame)
    assert profile.date_column == "date"
    assert profile.date_kind == "date_column"
    assert "segment" in profile.categorical_columns
    assert {"market_return", "smb_return"}.issubset(profile.return_columns)

    plan = plan_figure_suite(frame, title_prefix="Demo", max_figures=6)
    assert [item.kind for item in plan] == [
        "wide_time_series",
        "indexed_time_series",
        "cumulative_returns",
        "mean_bar",
        "distribution_by_category",
        "scatter_fit",
    ]

    with temp_figure_dir() as output_dir:
        result = create_figure_suite(
            frame,
            output_dir,
            style="ft",
            docx=True,
            source="",
            title_prefix="Demo",
            max_figures=6,
            formats=("png",),
        )
        assert not result.issues
        assert not result.skipped
        assert [figure.plan_item.kind for figure in result.generated_figures] == [
            "wide_time_series",
            "indexed_time_series",
            "cumulative_returns",
            "mean_bar",
            "distribution_by_category",
            "scatter_fit",
        ]
        assert result.docx_path is not None
        assert result.docx_path.exists()
        assert not validate_docx_images_fit_page(result.docx_path)
        proof_text = docx_text(result.docx_path)
        assert "FT-Style Figure Suite" in proof_text
        assert "Figure 1. Demo: Selected Series Over Time." in proof_text
        assert "The sample spans the time period 2018-01-31 to 2023." in proof_text
        assert "Source:" not in proof_text


def test_dataframe_figure_suite_detects_portfolio_return_narrative() -> None:
    rng = np.random.default_rng(2026)
    dates = pd.date_range("2015-01-31", periods=120, freq="ME")
    frame = pd.DataFrame(
        {
            "date": dates,
            "market_return": rng.normal(0.7, 3.7, len(dates)),
            "value_return": rng.normal(0.45, 2.7, len(dates)),
            "momentum_return": rng.normal(0.65, 4.2, len(dates)),
            "low_vol_return": rng.normal(0.35, 1.8, len(dates)),
        }
    )

    plan = plan_figure_suite(frame, title_prefix="Portfolio", narrative=True, max_figures=5)
    assert [item.kind for item in plan] == [
        "portfolio_growth_of_one_dollar",
        "portfolio_drawdowns",
        "portfolio_risk_return",
        "portfolio_return_correlations",
        "portfolio_return_distributions",
    ]

    with temp_figure_dir() as output_dir:
        result = create_figure_suite(
            frame,
            output_dir,
            style="ft",
            docx=True,
            source="",
            title_prefix="Portfolio",
            max_figures=5,
            narrative=True,
            formats=("png",),
        )
        assert not result.issues
        assert not result.skipped
        assert [figure.plan_item.kind for figure in result.generated_figures] == [
            "portfolio_growth_of_one_dollar",
            "portfolio_drawdowns",
            "portfolio_risk_return",
            "portfolio_return_correlations",
            "portfolio_return_distributions",
        ]
        assert result.docx_path is not None
        assert not validate_docx_images_fit_page(result.docx_path)
        proof_text = docx_text(result.docx_path)
        assert "Growth Of One Dollar Across Return Series" in proof_text
        assert "Risk And Return By Series" in proof_text
        assert "Annualized return and volatility" in proof_text
        assert "Source:" not in proof_text


def test_dataframe_figure_suite_detects_long_panel_composition() -> None:
    records = []
    for year in range(2010, 2025):
        for segment, base in [("All", 20.0), ("IG", 25.0), ("HY", 15.0)]:
            total = min(base + (year - 2010) * 2.2, 75.0)
            parts = {
                "MarketAxess": total if year < 2016 else total * 0.45,
                "Tradeweb": 0.0 if year < 2016 else total * 0.25,
                "Trumid": 0.0 if year < 2016 else total * 0.10,
                "Other electronic": 0.0 if year < 2016 else total * 0.20,
                "Total electronic": total,
                "Voice/high-touch": 100.0 - total,
            }
            for venue, share in parts.items():
                records.append(
                    {
                        "year": year,
                        "period": str(year),
                        "segment_group": segment,
                        "venue": venue,
                        "percentage_share": share,
                    }
                )
    frame = pd.DataFrame.from_records(records)

    profile = profile_dataframe(frame)
    assert profile.date_column == "year"
    assert profile.date_kind == "year_column"
    assert profile.sample == "2010 to 2024"

    plan = plan_figure_suite(frame, max_figures=8)
    kinds = [item.kind for item in plan]
    assert kinds[:4] == [
        "long_time_series",
        "long_part_to_whole",
        "long_slope",
        "long_latest_lollipop",
    ]
    assert plan[0].params["filter_column"] == "venue"
    assert plan[0].params["filter_value"] == "Total electronic"
    assert plan[1].params["segment"] == "venue"
    assert plan[1].params["filter_column"] == "segment_group"
    assert plan[1].params["filter_value"] == "All"

    narrative_plan = plan_figure_suite(frame, narrative=True, max_figures=5)
    assert [item.kind for item in narrative_plan] == [
        "electronic_voice_small_multiples",
        "electronic_whole_market_stacked_share",
        "electronic_venue_mix",
        "electronic_share_slope",
        "electronic_latest_venue_lollipop",
    ]
    extended_narrative_plan = plan_figure_suite(frame, narrative=True, max_figures=6)
    assert [item.kind for item in extended_narrative_plan][-1] == (
        "electronic_unconditional_bubble_matrix"
    )

    with temp_figure_dir() as output_dir:
        result = create_figure_suite(
            frame,
            output_dir,
            style="ft",
            docx=True,
            source="",
            title_prefix="Electronic Trading",
            max_figures=5,
            narrative=True,
            formats=("png",),
        )
        assert not result.issues
        assert not result.skipped
        assert [figure.plan_item.kind for figure in result.generated_figures] == [
            "electronic_voice_small_multiples",
            "electronic_whole_market_stacked_share",
            "electronic_venue_mix",
            "electronic_share_slope",
            "electronic_latest_venue_lollipop",
        ]
        assert result.docx_path is not None
        assert not validate_docx_images_fit_page(result.docx_path)
        proof_text = docx_text(result.docx_path)
        assert "Electronic Versus Voice Trading By Credit Segment" in proof_text
        assert "Whole-Market Trading Volume By Venue" in proof_text
        assert "Source:" not in proof_text

    with temp_figure_dir() as output_dir:
        result = create_figure_suite(
            frame,
            output_dir,
            style="ft",
            docx=True,
            source="",
            title_prefix="Electronic Trading",
            max_figures=6,
            narrative=True,
            formats=("png",),
        )
        assert not result.issues
        assert not result.skipped
        assert [figure.plan_item.kind for figure in result.generated_figures][-1] == (
            "electronic_unconditional_bubble_matrix"
        )
        assert result.docx_path is not None
        assert not validate_docx_images_fit_page(result.docx_path)


def test_validation_datasets_load_with_metadata() -> None:
    expected = {
        "ff3_monthly",
        "ff25_size_value_monthly",
        "ff_industry_10_monthly",
        "fred_financial_stress_daily",
        "fred_macro_monthly",
        "fred_rates_daily",
        "shiller_market_monthly",
        "world_bank_country_panel_annual",
        "world_bank_gdp_annual",
    }
    assert expected.issubset(set(available_validation_datasets()))

    dataset = load_validation_dataset("ff3_monthly")
    assert isinstance(dataset.data.index, pd.DatetimeIndex)
    assert dataset.data.index.min() <= pd.Timestamp("1927-01-31")
    assert {"Mkt-RF", "SMB", "HML", "RF"}.issubset(dataset.data.columns)
    assert dataset.units["Mkt-RF"] == "percent monthly return"
    assert "Kenneth French" in dataset.source

    macro = load_validation_dataset("fred_macro_monthly")
    assert {"UNRATE", "CPIAUCSL", "INDPRO", "PAYEMS", "FEDFUNDS"}.issubset(
        macro.data.columns
    )

    rates = load_validation_dataset("fred_rates_daily")
    assert {"DGS10", "DGS2", "DTB3", "T10Y2Y", "T10Y3M"}.issubset(rates.data.columns)

    shiller = load_validation_dataset("shiller_market_monthly")
    assert shiller.data.index.min() == pd.Timestamp("1871-01-31")
    assert {"real_price", "real_dividend", "real_earnings", "cape"}.issubset(
        shiller.data.columns
    )

    gdp = load_validation_dataset("world_bank_gdp_annual")
    assert {"country", "country_code", "gdp_trillions_usd"}.issubset(gdp.data.columns)
    assert gdp.data.index.min() == pd.Timestamp("2010-12-31")
    assert "World Bank" in gdp.source

    ff25 = load_validation_dataset("ff25_size_value_monthly")
    assert ff25.data.shape[1] == 25
    assert "SMALL LoBM" in ff25.data.columns
    assert "Kenneth French" in ff25.source

    stress = load_validation_dataset("fred_financial_stress_daily")
    assert {"VIXCLS", "BAMLH0A0HYM2"}.issubset(stress.data.columns)
    assert stress.data.index.min() <= pd.Timestamp("1990-01-02")

    panel = load_validation_dataset("world_bank_country_panel_annual")
    assert {"population_millions", "gdp_per_capita_usd"}.issubset(panel.data.columns)
    assert panel.data.index.min() == pd.Timestamp("2010-12-31")


def test_ft_theme_is_opt_in_and_keeps_background_optional() -> None:
    fins_rc = theme_rc()
    ft_rc = theme_rc(style="ft")
    ft_background_rc = theme_rc(style="ft", ft_background=True)

    assert fins_rc["figure.facecolor"] == "white"
    assert ft_rc["figure.facecolor"] == "white"
    assert ft_background_rc["figure.facecolor"] == FT_BACKGROUND
    assert ft_rc["axes.prop_cycle"] != fins_rc["axes.prop_cycle"]
    with pytest.raises(ValueError, match="style"):
        theme_rc(style="unknown")


def test_optional_plotly_ft_template_has_actionable_dependency_message() -> None:
    from fintools.figures.plotly_ft import ft_plotly_template

    try:
        import plotly  # noqa: F401
    except ImportError:
        with pytest.raises(RuntimeError, match="Optional Plotly figure support"):
            ft_plotly_template()
    else:
        template = ft_plotly_template()
        assert template.layout.paper_bgcolor == "white"


def test_core_plot_helpers_generate_labeled_axes() -> None:
    ff3 = load_validation_dataset("ff3_monthly").data
    ff25 = load_validation_dataset("ff25_size_value_monthly").data
    industries = load_validation_dataset("ff_industry_10_monthly").data
    macro = load_validation_dataset("fred_macro_monthly").data.reset_index()
    stress = load_validation_dataset("fred_financial_stress_daily").data
    gdp_panel = load_validation_dataset("world_bank_country_panel_annual").data.reset_index()
    gdp = load_validation_dataset("world_bank_gdp_annual").data.reset_index()

    fig, ax = time_series_plot(ff3, ["Mkt-RF", "SMB"], ylabel="Monthly return (%)")
    assert not validate_axes_labels(ax)
    assert not validate_horizontal_grid(ax)
    assert not validate_legend_present(ax)
    assert not validate_no_tick_label_overlap(ax)
    assert not validate_series_identification(ax)
    close(fig)

    fig, ax = time_series_plot(
        ff3,
        ["Mkt-RF", "SMB", "HML"],
        ylabel="Monthly return (%)",
        style="ft",
        direct_labels=True,
    )
    assert not validate_no_text_overlap(ax)
    assert not validate_legend_present(ax)
    close(fig)

    fig, ax = time_series_plot(ff3, "Mkt-RF", ylabel="Monthly return (%)")
    assert not validate_legend_present(ax)
    close(fig)

    fig, ax = cumulative_returns_plot(
        ff3,
        ["Mkt-RF", "SMB", "HML"],
        returns_are_percent=True,
        wealth_index=True,
        log_scale=True,
        style="ft",
        direct_labels=True,
    )
    assert ax.get_ylabel() == "Growth of $1"
    x_left, _ = ax.get_xlim()
    assert pd.Timestamp(plt.matplotlib.dates.num2date(x_left)).year <= 1927
    assert ax.patches
    tick_labels = [label.get_text() for label in ax.get_yticklabels()]
    assert any(label.startswith("$") for label in tick_labels)
    assert not any(label.endswith("x") for label in tick_labels)
    assert not validate_no_text_overlap(ax)
    labels = [text.get_text() for text in ax.get_legend().get_texts()]
    assert {"Mkt-RF", "SMB", "HML"}.issubset(labels)
    for line in ax.get_lines():
        assert line.get_alpha() == pytest.approx(0.78)
        assert line.get_linewidth() == pytest.approx(1.25)
    close(fig)

    fig, ax = indexed_time_series_plot(
        load_validation_dataset("fred_macro_monthly").data,
        ["INDPRO", "PAYEMS", "CPIAUCSL"],
    )
    assert ax.get_ylabel().startswith("Index")
    assert not validate_no_tick_label_overlap(ax)
    assert not validate_category_label_count(ax)
    close(fig)

    fig, ax = drawdown_plot(ff3, "Mkt-RF", returns_are_percent=True)
    assert ax.get_ylabel() == "Drawdown"
    close(fig)

    means = industries.mean().reset_index()
    means.columns = ["industry", "return"]
    fig, ax = bar_plot(means, "industry", "return", ylabel="Mean monthly return (%)")
    assert not validate_axes_labels(ax)
    close(fig)

    fig, ax, summary = mean_return_bar_plot(industries, error="se")
    assert not validate_axes_labels(ax)
    assert "se" in summary.columns
    assert not any("Error bars:" in text.get_text() for text in ax.texts)
    bar_colors = {patch.get_facecolor() for patch in ax.patches}
    assert len(bar_colors) > 1
    close(fig)

    fig, ax = stacked_bar_plot(ff3, ["Mkt-RF", "SMB", "HML"], max_bars=24, style="ft")
    assert not validate_axes_labels(ax)
    assert ax.containers
    assert all(patch.get_alpha() == pytest.approx(0.72) for patch in ax.patches)
    assert len([label for label in ax.get_xticklabels() if label.get_text()]) <= 8
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    long_macro = macro.melt(id_vars="date", value_vars=["UNRATE", "FEDFUNDS"])
    fig, ax = grouped_bar_plot(
        long_macro.tail(12),
        "date",
        "value",
        "variable",
        ylabel="Percent",
    )
    assert not validate_axes_labels(ax)
    close(fig)

    fig, ax = scatter_plot(macro, "FEDFUNDS", "UNRATE", fit=True)
    assert not validate_axes_labels(ax)
    assert any("R-squared" in text.get_text() for text in ax.texts)
    close(fig)

    fig, ax = distribution_plot(ff3.reset_index(), "Mkt-RF")
    assert not validate_axes_labels(ax)
    close(fig)

    fig, ax = correlation_heatmap(industries)
    assert ax.get_title(loc="left")
    assert not validate_no_tick_label_overlap(ax)
    assert not validate_no_tick_label_overlap(ax, axis="y")
    close(fig)

    fred_frame = pd.DataFrame(
        np.random.default_rng(11).normal(size=(80, 8)),
        columns=[
            "DGS10",
            "DGS2",
            "DTB3",
            "T10Y2Y",
            "VIXCLS",
            "BAMLH0A0HYM2",
            "ten_year_minus_two_year",
            "ten_year_minus_three_month",
        ],
    )
    fig, ax = correlation_heatmap(
        fred_frame,
        profile="word_a4",
        style="ft",
        xlabel="Series",
        ylabel="Series",
    )
    x_labels = [label.get_text() for label in ax.get_xticklabels()]
    assert "10Y-3M" in x_labels
    assert "HY OAS" in x_labels
    assert ax.get_xlabel() == "Series"
    assert ax.get_ylabel() == "Series"
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    fig, axes = small_multiples(
        industries,
        ["NoDur", "Durbl", "HiTec", "Utils"],
        ylabel=["Nondurables", "Durables", "High tech", "Utilities"],
        style="ft",
    )
    assert len(axes) >= 4
    for ax in axes[:4]:
        assert ax.get_title(loc="left")
        assert ax.get_title(loc="left") in {"NoDur", "Durbl", "HiTec", "Utils"}
        assert ax.lines[0].get_linewidth() == pytest.approx(1.0)
        assert ax.lines[0].get_alpha() == pytest.approx(0.82)
        assert not validate_no_tick_label_overlap(ax)
        assert all(label.get_rotation() == 0 for label in ax.get_xticklabels())
    assert [ax.get_ylabel() for ax in axes[:4]] == [
        "Nondurables",
        "Durables",
        "High tech",
        "Utilities",
    ]
    close(fig)

    ff25_mean = pd.DataFrame(
        [
            {
                "size": ["Small", "2", "3", "4", "Big"][index // 5],
                "value": ["Low BM", "2", "3", "4", "High BM"][index % 5],
                "mean_return": ff25[column].mean(),
            }
            for index, column in enumerate(ff25.columns)
        ]
    )
    fig, ax = value_heatmap(ff25_mean, "size", "value", "mean_return", style="ft")
    assert ax.get_title(loc="left")
    assert ax.get_ylabel() == "Size"
    assert ax.get_xlabel() == "Value"
    assert [label.get_text() for label in ax.get_xticklabels()] == [
        "Low BM",
        "2",
        "3",
        "4",
        "High BM",
    ]
    assert not validate_display_labels(ax)
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    ff25_long = pd.concat(
        [
            pd.DataFrame(
                {
                    "return": ff25[column],
                    "value": ["Low BM", "2", "3", "4", "High BM"][index % 5],
                }
            )
            for index, column in enumerate(ff25.columns)
        ],
        ignore_index=True,
    )
    fig, ax = distribution_comparison_plot(
        ff25_long,
        "return",
        "value",
        kind="box",
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert ax.get_xlabel() == "Value"
    assert not validate_display_labels(ax)
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    industry_excess = (
        industries.mean()
        .sub(float(ff3["Mkt-RF"].mean()))
        .rename("spread")
        .reset_index()
        .rename(columns={"index": "industry"})
    )
    fig, ax = diverging_bar_plot(industry_excess, "industry", "spread", style="ft")
    assert not validate_axes_labels(ax)
    assert not validate_display_labels(ax)
    assert len({patch.get_facecolor() for patch in ax.patches}) > 1
    close(fig)

    fig, ax = ecdf_plot(industries, ["NoDur", "HiTec", "Enrgy"], style="ft")
    assert not validate_axes_labels(ax)
    assert not validate_legend_present(ax)
    close(fig)

    latest_gdp = gdp[gdp["date"].dt.year == 2024]
    fig, ax = lollipop_plot(
        latest_gdp,
        "country",
        "gdp_trillions_usd",
        xlabel="GDP (trillions of current U.S. dollars)",
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert ax.get_ylabel() == "Country"
    assert not validate_display_labels(ax)
    close(fig)

    fig, ax = lollipop_plot(
        latest_gdp,
        "country",
        "gdp_trillions_usd",
        highlight=["United States", "China", "India"],
        style="ft",
    )
    assert not validate_display_labels(ax)
    point_colors = {
        tuple(facecolor)
        for collection in ax.collections
        for facecolor in collection.get_facecolors()
    }
    assert len(point_colors) >= 4
    close(fig)

    gdp_panel["year"] = gdp_panel["date"].dt.year
    gdp_per_capita = (
        gdp_panel[gdp_panel["year"].isin([2010, 2024])]
        .pivot(index="country", columns="year", values="gdp_per_capita_usd")
        .reset_index()
        .rename(columns={2010: "gdp_pc_2010", 2024: "gdp_pc_2024"})
    )
    fig, ax = dumbbell_plot(
        gdp_per_capita,
        "country",
        "gdp_pc_2010",
        "gdp_pc_2024",
        limit=8,
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert ax.get_ylabel() == "Country"
    assert not validate_display_labels(ax)
    assert not validate_legend_present(ax)
    close(fig)

    latest_panel = gdp_panel[gdp_panel["year"] == 2024].copy()
    fig, ax = bubble_scatter_plot(
        latest_panel,
        "population_millions",
        "gdp_per_capita_usd",
        "gdp_trillions_usd",
        label="country",
        label_top=3,
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert not validate_display_labels(ax)
    assert not validate_no_text_overlap(ax)
    assert not validate_markers_within_axes(ax)
    assert ax.get_ylim()[1] > latest_panel["gdp_per_capita_usd"].max()
    close(fig)

    share_frame = gdp_panel[
        gdp_panel["year"].isin([2010, 2024])
        & gdp_panel["country"].isin(["United States", "China", "Japan"])
    ].copy()
    fig, ax = proportional_stacked_bar_plot(
        share_frame,
        "year",
        "country",
        "gdp_current_usd",
        colors={"United States": "#111111", "China": "#222222", "Japan": "#333333"},
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert ax.get_xlabel() == "Year"
    assert [label.get_text() for label in ax.get_xticklabels()] == ["2010", "2024"]
    assert ax.get_legend() is not None
    assert ax.get_legend().get_title().get_text() == ""
    assert {patch.get_facecolor() for patch in ax.patches}
    assert not validate_display_labels(ax)
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    area_frame = pd.DataFrame(
        {
            "year": [2020, 2021, 2022],
            "Electronic": [0.30, 0.36, 0.42],
            "Voice/high-touch": [0.70, 0.64, 0.58],
        }
    )
    fig, ax = stacked_area_plot(
        area_frame,
        "year",
        ["Electronic", "Voice/high-touch"],
        colors={"Electronic": "#111111", "Voice/high-touch": "#999999"},
        style="ft",
    )
    assert ax.get_xlabel() == "Year"
    assert ax.get_ylabel() == "Share"
    assert ax.get_ylim()[1] == pytest.approx(1.0)
    assert ax.get_legend() is not None
    assert ax.get_legend().get_title().get_text() == ""
    assert not validate_display_labels(ax)
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    slope_frame = (
        gdp.assign(year=gdp["date"].dt.year)
        .pivot(index="country", columns="year", values="gdp_trillions_usd")
        .reset_index()
        .rename(columns={2010: "gdp_2010", 2024: "gdp_2024"})
    )
    fig, ax = slope_chart(
        slope_frame,
        "country",
        "gdp_2010",
        "gdp_2024",
        ylabel="GDP (trillions of current U.S. dollars)",
        limit=8,
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert not validate_no_text_overlap(ax)
    assert not validate_unique_series_colors(ax, minimum=8)
    close(fig)

    macro_indexed = load_validation_dataset("fred_macro_monthly").data
    fig, ax = connected_scatter_plot(
        macro_indexed[["FEDFUNDS", "UNRATE"]].dropna().loc["2007":"2012"],
        "FEDFUNDS",
        "UNRATE",
        xlabel="Federal funds rate (%)",
        ylabel="Unemployment rate (%)",
        max_points=24,
        style="ft",
    )
    assert not validate_axes_labels(ax)
    close(fig)

    with pytest.raises(ValueError, match="short episodes"):
        connected_scatter_plot(
            macro_indexed[["FEDFUNDS", "UNRATE"]].dropna().loc["1990":],
            "FEDFUNDS",
            "UNRATE",
            style="ft",
        )

    fig, ax = area_balance_plot(
        load_validation_dataset("fred_rates_daily").data[["T10Y2Y"]].dropna(),
        "T10Y2Y",
        ylabel="Percentage points",
        style="ft",
    )
    assert ax.patches
    assert not validate_legend_present(ax)
    close(fig)

    fig, ax = calendar_heatmap(stress, "VIXCLS", year=2020, style="ft")
    assert not validate_axes_labels(ax)
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    fig, ax = rolling_stat_plot(
        stress,
        "VIXCLS",
        window=21,
        statistic="volatility",
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert not validate_no_tick_label_overlap(ax)
    close(fig)

    vix = stress["VIXCLS"].dropna()
    rolling_mean = vix.rolling(window=63, min_periods=21).mean()
    rolling_std = vix.rolling(window=63, min_periods=21).std()
    band_frame = pd.DataFrame(
        {
            "rolling_mean": rolling_mean,
            "lower": (rolling_mean - rolling_std).clip(lower=0),
            "upper": rolling_mean + rolling_std,
        }
    ).dropna()
    fig, ax = uncertainty_band_plot(
        band_frame,
        "rolling_mean",
        "lower",
        "upper",
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert not validate_legend_present(ax)
    close(fig)


def test_nber_recession_windows_are_clipped_to_plot_range() -> None:
    windows = recession_windows_for_range("2020-01-31", "2021-12-31")
    assert windows == [
        (
            pd.Timestamp("2020-03-01"),
            pd.Timestamp("2020-04-30"),
            "2020",
        )
    ]

    fig, ax = plt.subplots()
    ax.plot([pd.Timestamp("2020-01-31"), pd.Timestamp("2020-12-31")], [0, 1])
    artists = add_nber_recession_shading(
        ax,
        data_start=pd.Timestamp("2020-01-31"),
        data_end=pd.Timestamp("2020-12-31"),
    )
    assert artists
    assert all(getattr(artist, "get_hatch", lambda: None)() in (None, "") for artist in artists)
    facecolor = artists[0].get_facecolor()
    assert facecolor[0] > 0.85 and facecolor[1] > 0.85 and facecolor[2] > 0.85
    assert facecolor[3] >= 0.55
    assert artists[0].get_zorder() >= 0
    close(fig)


def test_bubble_matrix_plot_and_layout_validators() -> None:
    frame = pd.DataFrame(
        {
            "venue": ["MarketAxess", "Tradeweb", "Voice", "MarketAxess", "Tradeweb", "Voice"],
            "segment": ["IG", "IG", "IG", "HY", "HY", "HY"],
            "share": [18.0, 9.0, 73.0, 9.0, 4.0, 87.0],
        }
    )
    fig, ax = bubble_matrix_plot(
        frame,
        "venue",
        "segment",
        "share",
        title="Average Venue Shares",
        xlabel="Venue",
        ylabel="Segment",
        size_label="Average share",
        x_order=["MarketAxess", "Tradeweb", "Voice"],
        y_order=["IG", "HY"],
        colors={"MarketAxess": "#1F77B4", "Tradeweb": "#3A9A44", "Voice": "#9A3A4D"},
        style="ft",
    )
    assert not validate_axes_labels(ax)
    assert not validate_display_labels(ax)
    assert not validate_markers_within_axes(ax)
    assert not validate_no_tick_label_overlap(ax)
    assert not validate_no_tick_label_overlap(ax, axis="y")
    assert not validate_titles_within_canvas(fig)
    assert ax.get_legend() is not None
    close(fig)

    fig, axes = plt.subplots(ncols=2)
    axes[0].plot([0, 1], [0, 1])
    axes[1].plot([0, 1], [1, 0])
    assert not validate_equal_subplot_widths(fig)
    axes[1].set_position([0.70, 0.11, 0.15, 0.77])
    issues = validate_equal_subplot_widths(fig)
    assert issues and issues[0].code == "unequal_subplot_widths"
    close(fig)

    fig, ax = plt.subplots()
    ax.plot([0, 1], [0, 1])
    fig.suptitle("Outside", x=1.25)
    issues = validate_titles_within_canvas(fig)
    assert issues and issues[0].code == "title_outside_canvas"
    close(fig)


def test_export_bundle_and_word_docx() -> None:
    ff3 = load_validation_dataset("ff3_monthly").data
    context = FigureContext(
        title="Cumulative Market Excess Return",
        note="Validation figure for Word/A4 export.",
        source="Kenneth French Data Library",
        sample="2020-01-31 to 2021-12-31",
        units="Cumulative return",
    )
    assert context.caption_text(3).startswith("Figure 3. Cumulative Market Excess Return.")
    assert (
        "The sample spans the time period 2020-01-31 to 2021-12-31."
        in context.caption_text(3)
    )
    fig, _ = cumulative_returns_plot(ff3, "Mkt-RF", profile="word_a4")

    with temp_figure_dir() as output_dir:
        bundle = export_figure_bundle(fig, output_dir, "market_cumulative", context=context)
        assert bundle["png"].exists() and bundle["png"].stat().st_size > 0
        assert bundle["pdf"].exists() and bundle["pdf"].stat().st_size > 0
        assert "Kenneth French" in bundle["caption"].read_text(encoding="utf-8")
        assert not validate_image_not_blank(bundle["png"])

        word_bundle = export_word_figure(
            fig,
            output_dir,
            "market_cumulative_word",
            context=context,
        )
        assert word_bundle["png"].exists() and word_bundle["png"].stat().st_size > 0
        assert not validate_figure_context(context)
        assert not validate_word_readability(fig)

        docx_path = insert_figure_docx(
            word_bundle["png"],
            output_dir / "figure_sheet.docx",
            context=context,
            spec="landscape_wide",
        )
        assert docx_path.exists() and docx_path.stat().st_size > 0
        assert not validate_docx_images_fit_page(docx_path)
        single_text = docx_text(docx_path)
        assert "Figure 1. Cumulative Market Excess Return." in single_text
        assert "Validation figure for Word/A4 export." in single_text
        assert "The sample spans the time period 2020-01-31 to 2021-12-31." in single_text
        assert "Units: Cumulative return." in single_text
        assert "Source: Kenneth French Data Library." in single_text

        combined_docx_path = insert_figures_docx(
            [
                WordFigureEntry(word_bundle["png"], context=context),
                WordFigureEntry(word_bundle["png"], context=context, spec="landscape_wide"),
            ],
            output_dir / "figure_pack.docx",
            title="Validation Figure Pack",
        )
        assert combined_docx_path.exists() and combined_docx_path.stat().st_size > 0
        assert not validate_docx_images_fit_page(combined_docx_path)
        combined_text = docx_text(combined_docx_path)
        assert "Validation Figure Pack" in combined_text
        assert "Figure 1. Cumulative Market Excess Return." in combined_text
        assert "Figure 2. Cumulative Market Excess Return." in combined_text
        close(fig)


def test_return_scale_inference_flags_common_units() -> None:
    assert infer_return_scale([1.0, -2.5, 3.0]) == "percent"
    assert infer_return_scale([0.01, -0.02, 0.03]) == "decimal"


def test_display_label_validation_flags_raw_field_names() -> None:
    fig, ax = plt.subplots()
    ax.set_xlabel("country")
    ax.set_ylabel("gdp_current_usd")
    issues = validate_display_labels(ax)
    assert {issue.code for issue in issues} == {"raw_display_label"}
    close(fig)

    fig, ax = plt.subplots()
    ax.set_xlabel("Country")
    ax.set_ylabel("GDP current USD")
    assert not validate_display_labels(ax)
    close(fig)


def test_value_heatmap_suppresses_visual_negative_zero() -> None:
    frame = pd.DataFrame({"row": ["A"], "column": ["B"], "value": [-0.0001]})

    fig, ax = value_heatmap(frame, "row", "column", "value", fmt=".1f")

    assert [text.get_text() for text in ax.texts] == ["0.0"]
    close(fig)


def test_add_source_note_places_standalone_source_text() -> None:
    fig, _ax = plt.subplots()

    note = add_source_note(fig, "Source: validation dataset.", style="ft")

    assert note.get_text() == "Source: validation dataset."
    assert note.figure is fig
    close(fig)


def test_marker_validation_flags_clipped_bubbles() -> None:
    fig, ax = plt.subplots()
    ax.scatter([1.0], [1.0], s=[900.0])
    ax.set_xlim(0.0, 1.0)
    ax.set_ylim(0.0, 1.0)
    issues = validate_markers_within_axes(ax)
    assert issues and issues[0].code == "marker_clipped"
    close(fig)


def test_build_figure_workflow_records_ft_references() -> None:
    text = (ROOT / "docs" / "ai" / "workflows" / "build-figure.md").read_text(
        encoding="utf-8"
    )
    assert "github.com/Financial-Times/chart-doctor" in text
    assert "style-your-visuals-like-the-financial-times-using-plotly" in text
    assert "aeturrell.github.io/coding-for-economists/vis-narrative.html" in text
